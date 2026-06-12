"""Generate PM Studio AI Action Jobs reference DOCX."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parents[2] / "docs" / "PM-Studio-AI-Action-Jobs-Reference.docx"

AI_JOBS = [
    ("1", "Requirements", "Upload requirement PDF (auto)", "req_analyze",
     "Reads extracted PDF text; classifies gaps (client vs technical), risks, project type, auto-answers technical gaps.",
     "Foundation of entire project — wrong gaps lead to wrong PRD/SRS.",
     "Long-context document analysis + structured JSON extraction"),
    ("2", "Requirements", "Regenerate analysis", "req_analyze",
     "Re-runs same analysis on existing PDF.",
     "Fixes bad first pass without re-upload.",
     "Long-context document analysis + structured JSON extraction"),
    ("3", "Requirements", "Synthesize feedback", "req_synthesize",
     "Merges original requirement + client feedback into final draft (features, scope, decisions).",
     "Locks scope before PRD; reduces client disputes.",
     "Cross-document synthesis + business requirements writing"),
    ("4", "Requirements", "Rewrite analysis (PM instructions)", "req_synthesize",
     "Rewrites final draft per PM edit instructions.",
     "Iterates scope without new upload.",
     "Instruction-following rewrite + structured JSON"),
    ("5", "PRDs", "Generate PRD", "prd_generate",
     "Builds full PRD: personas, features, user stories, acceptance criteria, metrics.",
     "Commercial contract with client; drives all downstream docs.",
     "Product/business narrative + nested structured JSON"),
    ("6", "PRDs", "Regenerate PRD", "prd_generate",
     "Rebuilds PRD from requirement analysis.",
     "Recovery from weak generation.",
     "Product/business narrative + nested structured JSON"),
    ("7", "PRD detail", "Rewrite PRD (PM comment)", "prd_rewrite",
     "Rewrites existing PRD sections per PM instructions.",
     "Refinement without full regen.",
     "Instruction-following rewrite + product writing"),
    ("8", "SRS", "Generate SRS", "srs_generate",
     "IEEE-style SRS: FR-001…, NFRs, data entities, interfaces.",
     "Legal/technical spec — every later task links to FR numbers.",
     "Formal standards compliance (IEEE 830) + precise structured JSON"),
    ("9", "SRS", "Regenerate SRS", "srs_generate",
     "Rebuilds SRS from PRD.",
     "Fixes missing FRs or weak NFRs.",
     "Formal standards compliance (IEEE 830) + precise structured JSON"),
    ("10", "SRS detail", "Rewrite SRS (PM comment)", "srs_rewrite",
     "Rewrites SRS while preserving FR numbering.",
     "Surgical fixes; FR IDs must not break.",
     "Instruction-following rewrite + traceability preservation"),
    ("11", "Architecture", "Generate Architecture Suite", "arch_generate",
     "Generates 6 docs: System, DB, API, Frontend, Security, UI/UX (large JSON + Mermaid).",
     "Build blueprint for dev team; API paths, tables, auth model.",
     "Systems architecture reasoning + very long context + structured JSON"),
    ("12", "Architecture", "Resume generation", "arch_generate",
     "Continues suite from first incomplete doc.",
     "Completes partial runs.",
     "Systems architecture reasoning + very long context + structured JSON"),
    ("13", "Architecture", "Regenerate full suite", "arch_generate",
     "Clears all 6 docs and regenerates.",
     "Nuclear reset when suite is misaligned.",
     "Systems architecture reasoning + very long context + structured JSON"),
    ("14", "Architecture detail", "Generate single doc", "arch_generate",
     "One of 6 architecture documents.",
     "Targeted fix without burning quota on all 6.",
     "Domain-specific architecture (DB/API/FE/security)"),
    ("15", "Architecture detail", "Regenerate single doc", "arch_generate",
     "Replaces one doc.",
     "Refresh weak doc (e.g. API only).",
     "Domain-specific architecture (DB/API/FE/security)"),
    ("16", "Architecture detail", "Apply AI / AI-assist", "arch_generate",
     "Regenerates one doc incorporating PM instructions.",
     "Human-guided correction.",
     "Instruction-following + architecture reasoning"),
    ("17", "Tasks / Kanban", "Generate tasks from SRS & Architecture", "module_extract",
     "Maps every FR to Kanban tasks with file paths, endpoints, tables from architecture.",
     "FR to implementable work; gaps here = missed features.",
     "Entity-to-structure mapping + software project planning"),
    ("18", "Tasks", "Fill gaps / replace modes", "module_extract",
     "Only missing FRs, or wipe + regenerate tasks.",
     "Coverage repair without duplicating tasks.",
     "Entity-to-structure mapping + software project planning"),
    ("19", "Task detail", "Generate technical spec", "spec_generate",
     "Full dev spec: files, steps, DB, API, security, Cursor prompt.",
     "Direct input to Cursor/IDE — most code-quality-critical job.",
     "Software engineering / coding-agent + zero-ambiguity spec writing"),
    ("20", "Task detail", "Regenerate spec", "spec_generate",
     "Replaces failed or weak spec.",
     "Retry without new task.",
     "Software engineering / coding-agent + zero-ambiguity spec writing"),
    ("21", "Tasks", "Generate project orchestration", "orchestration_generate",
     "Master build guide: file manifest, build order, Cursor workspace prompt.",
     "Single paste-in prompt for whole project.",
     "Cross-document aggregation + implementation sequencing"),
]

NON_AI = [
    ("Requirements", "Cost estimate", "Local math from feature count — no LLM"),
    ("PRD / SRS detail", "Quality check", "Rule-based checklist scoring — no LLM"),
    ("PRD / SRS detail", "Export PDF", "Document export — no LLM"),
    ("Architecture detail", "Align suite", "Deterministic cross-doc fixes — no LLM"),
    ("Architecture", "Save / Edit JSON", "Manual edit — no LLM"),
    ("Knowledge, Decisions, Clients, Projects", "All actions", "No AI in current codebase"),
]

SPECIALTIES = [
    ("Long-context document analysis", "#1, #2", "Reads full PDF; does not miss sections"),
    ("Structured JSON / schema adherence", "All 21 jobs", "Valid Pydantic output every time"),
    ("Business / product narrative", "#5, #6, #7", "Client-readable PRD, clear priorities"),
    ("Formal spec / IEEE compliance", "#8, #9, #10", "FR-001… numbering, measurable NFRs"),
    ("Systems architecture reasoning", "#11–#16", "Consistent auth, API, DB, monolith signals"),
    ("Cross-document synthesis", "#3, #4, #21", "Merges sources without contradiction"),
    ("Instruction-following rewrite", "#4, #7, #10, #16", "Changes only what PM asked; preserves IDs"),
    ("Entity-to-structure mapping", "#17, #18", "Every FR → task + real file path"),
    ("Software engineering / coding-agent", "#19, #20", "Exact paths, functions, Cursor-ready prompts"),
    ("Implementation sequencing", "#21", "Correct build order across 50+ files"),
]

GAPS = [
    ("orchestration_generate not in FREE_ROUTING", "Job #21 fails unless Anthropic key exists"),
    ("arch_single_doc routing unused", "Code always uses arch_generate — same chain in practice"),
    ("quality_check / cost_estimate in routing", "Defined but never called by UI"),
    ("Per-screen model override", "Disables 14-model fallback for that screen's jobs"),
]


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def add_table(doc: Document, headers: list[str], rows: list[tuple], col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=9)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            set_cell_text(cells[ci], str(val), size=8)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Title
    title = doc.add_heading("PM Studio — AI Action Jobs Reference", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Generated {date.today().isoformat()} · PM Studio codebase audit")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_heading("Pipeline overview", level=1)
    doc.add_paragraph(
        "Upload PDF → Analyze → Synthesize feedback → PRD → SRS → "
        "Architecture (6 docs) → Kanban tasks → Per-task specs → Orchestration"
    )
    doc.add_paragraph(
        "Each step feeds the next. Weak output upstream propagates downstream."
    )

    doc.add_heading("Master table — all 21 AI action jobs", level=1)
    add_table(
        doc,
        ["#", "Screen", "UI action", "task_type", "What it does", "Why important", "Ideal LLM specialty"],
        AI_JOBS,
        col_widths=[0.35, 0.85, 1.1, 0.95, 1.5, 1.1, 1.15],
    )

    doc.add_heading("Buttons that look like AI but are NOT LLM jobs", level=1)
    add_table(doc, ["Screen", "Button", "Actual behavior"], NON_AI, [1.2, 1.5, 3.3])

    doc.add_heading("LLM specialty types — quick reference", level=1)
    add_table(doc, ["Specialty type", "Job numbers", "What good looks like"], SPECIALTIES, [1.8, 0.8, 3.4])

    doc.add_heading("Free mode — configured primary models (live org)", level=1)
    add_table(
        doc,
        ["Screen", "Primary model", "Specialty fit"],
        [
            ("Requirements", "NVIDIA Nemotron 3 Super (FREE)", "Long-context"),
            ("PRD / SRS", "OpenAI GPT-OSS 120B (FREE)", "Structured JSON"),
            ("Architecture", "NVIDIA Nemotron Ultra 550B (FREE)", "Largest reasoning / architecture"),
            ("Tasks (specs)", "Poolside Laguna M.1 (FREE)", "Coding-agent"),
            ("Tasks (extract)", "Kimi K2.6 (FREE)", "Structured extraction"),
        ],
        [1.2, 2.2, 2.6],
    )

    doc.add_heading("Fallback behavior (free mode)", level=1)
    for bullet in [
        "Up to 14 models per job, tried in order automatically.",
        "Rate limit (429): wait 1 second, then next model.",
        "Dead model / timeout (12 min on arch jobs) / API error: skip immediately to next.",
        "Providers: OpenRouter (primary) → Groq → Together AI (if key configured).",
        "After all free models fail: Anthropic Claude Sonnet if key exists.",
        "Per-screen model override disables the entire fallback chain.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Priority ranking for quality investment", level=1)
    priorities = [
        ("1", "#19–#20 Task spec", "Coding-agent specialty — feeds Cursor directly"),
        ("2", "#8–#10 SRS", "Formal spec + traceability specialty"),
        ("3", "#11–#16 Architecture", "Systems architecture + long context"),
        ("4", "#17–#18 Module extract", "Mapping specialty (FR coverage)"),
        ("5", "#5–#7 PRD", "Product writing specialty"),
        ("6", "#1–#4 Requirements", "Document analysis specialty"),
        ("7", "#21 Orchestration", "Aggregation specialty (broken in free mode today)"),
    ]
    add_table(doc, ["Priority", "Jobs", "Reason"], priorities, [0.6, 1.4, 4.0])

    doc.add_heading("Known gaps (codebase audit)", level=1)
    add_table(doc, ["Issue", "Impact"], GAPS, [2.5, 3.5])

    doc.add_paragraph()
    p = doc.add_paragraph("Total: 21 distinct AI action patterns across 5 screens.")
    p.runs[0].bold = True

    doc.save(OUTPUT)
    print(f"Written: {OUTPUT}")


if __name__ == "__main__":
    main()
